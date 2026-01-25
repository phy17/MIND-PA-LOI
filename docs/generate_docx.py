from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 标题
title = doc.add_heading('E.R.A-MIND 论文创新点总结', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 创新点 1
doc.add_heading('创新点 1：风险感知规划（安全来源）', level=1)

doc.add_heading('来源论文', level=2)
doc.add_paragraph('MARC: Multipolicy and Risk-aware Contingency Planning for Autonomous Driving')
doc.add_paragraph('发表于 IEEE Robotics and Automation Letters (RA-L), 2023')

doc.add_heading('核心思想', level=2)
doc.add_paragraph('针对"鬼探头"等高风险低概率场景，引入基于马氏距离的风险感知规划机制。')

doc.add_heading('关键公式', level=2)
doc.add_paragraph('来自 MARC 论文 Eq. 12 和 Eq. 14：')
doc.add_paragraph('Safety Cost: l_t = l_t^safe + l_t^tar + l_t^kin + l_t^comf')
doc.add_paragraph('l_t^safe = Σ G(max(D_bnd - D(N_t^j), 0))')
doc.add_paragraph('其中：')
doc.add_paragraph('• D(N_t^j) 是 Ego 车辆到障碍物（高斯分布）的马氏距离')
doc.add_paragraph('• D_bnd 是安全边界距离')
doc.add_paragraph('• 当距离小于安全边界时，惩罚项 G 会急剧增加')

doc.add_heading('实施方式', level=2)
doc.add_paragraph('将马氏距离风险计算集成到 MIND 的场景树剪枝逻辑和 iLQR 轨迹优化中，确保规划器不会忽略低概率但高风险的场景。')

# 分隔线
doc.add_paragraph('─' * 50)

# 创新点 2
doc.add_heading('创新点 2：实例中心场景表示（效率来源）', level=1)

doc.add_heading('来源论文', level=2)
doc.add_paragraph('SIMPL: A Simple and Efficient Multi-agent Motion Prediction Baseline for Autonomous Driving')
doc.add_paragraph('发表于 IEEE Robotics and Automation Letters (RA-L), 2024')

doc.add_heading('核心思想', level=2)
doc.add_paragraph('采用实例中心（Instance-centric）的场景表示方法，为每个交通参与者建立局部坐标系，提高计算效率和预测精度。')

doc.add_heading('关键公式', level=2)
doc.add_paragraph('来自 SIMPL 论文 Section III.C：')
doc.add_paragraph('相对位姿编码: sin(α_i→j) = (v_i × v_j) / (||v_i|| ||v_j||)')
doc.add_paragraph('其中：')
doc.add_paragraph('• α_i→j 是车辆 i 相对于车辆 j 的朝向角')
doc.add_paragraph('• v_i, v_j 是两车的速度向量')

doc.add_heading('实施方式', level=2)
doc.add_paragraph('复用 SIMPL 的相对位姿计算模块来评估场景复杂度（Scene Chaos Index）。当周围车辆的相对朝向分布混乱时，说明路况复杂，系统自动开启"防御模式"。')

# 分隔线
doc.add_paragraph('─' * 50)

# 创新点 3
doc.add_heading('创新点 3：流效率规划（流效率来源）', level=1)

doc.add_heading('来源论文', level=2)
doc.add_paragraph('EPSILON: An Efficient Planning System for Automated Vehicles in Highly Interactive Environments')
doc.add_paragraph('发表于 IEEE Transactions on Robotics (T-RO), 2021')

doc.add_heading('核心思想', level=2)
doc.add_paragraph('引入多维度的流效率评估机制，使车辆在跟车场景中表现得更像"老司机"——该快快，该慢慢。')

doc.add_heading('关键公式', level=2)
doc.add_paragraph('来自 EPSILON 论文 Section VII.D：')
doc.add_paragraph('流效率函数: F_e = Σ(λ_e^p Δv_p + λ_e^o Δv_o + λ_e^l Δv_l)')
doc.add_paragraph('其中：')
doc.add_paragraph('• Δv_p = |v_ego - v_pref|：个人效率损失（我和期望速度的差距）')
doc.add_paragraph('• Δv_l = |v_lead - v_pref|：环境限制损失（前车对我的速度压制）')
doc.add_paragraph('• Δv_o = max(v_ego - v_lead, 0)：超调风险（我是否比前车开得太快）')

doc.add_heading('实施方式', level=2)
doc.add_paragraph('将 EPSILON 的三项流效率指标（个人欲望、前车压制、超速惩罚）集成到 MIND 的代价函数中，替代原有的简单 TargetSpeed 目标。')

# 分隔线
doc.add_paragraph('─' * 50)

# 总结表格
doc.add_heading('总结', level=1)

table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'

# 表头
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '创新点'
hdr_cells[1].text = '来源论文'
hdr_cells[2].text = '核心贡献'

# 数据行
row1 = table.rows[1].cells
row1[0].text = '风险感知规划'
row1[1].text = 'MARC (RA-L 2023)'
row1[2].text = '基于马氏距离的安全代价，应对鬼探头等高风险场景'

row2 = table.rows[2].cells
row2[0].text = '实例中心表示'
row2[1].text = 'SIMPL (RA-L 2024)'
row2[2].text = '相对位姿编码，提升计算效率，评估场景复杂度'

row3 = table.rows[3].cells
row3[0].text = '流效率规划'
row3[1].text = 'EPSILON (T-RO 2021)'
row3[2].text = '多维度速度代价，实现智能跟车行为'

# 参考文献
doc.add_heading('参考文献', level=1)
doc.add_paragraph('1. Tong, J., et al. "MARC: Multipolicy and Risk-aware Contingency Planning for Autonomous Driving." IEEE Robotics and Automation Letters, 2023.')
doc.add_paragraph('2. Zhang, L., et al. "SIMPL: A Simple and Efficient Multi-agent Motion Prediction Baseline for Autonomous Driving." IEEE Robotics and Automation Letters, 2024.')
doc.add_paragraph('3. Luo, W., et al. "EPSILON: An Efficient Planning System for Automated Vehicles in Highly Interactive Environments." IEEE Transactions on Robotics, 2021.')

# 保存
doc.save('/Users/phy/Desktop/MIND/docs/MIND_Improvements_Summary.docx')
print('Word 文档已生成！')
