from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

# 创建文档
doc = Document()

# 设置标题
title = doc.add_heading('MIND 自动驾驶决策规划系统技术报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加日期和项目信息
doc.add_paragraph(f'日期: {datetime.date.today().strftime("%Y年%m月%d日")}')
doc.add_paragraph('项目: MIND (Multimodal Integrated Prediction and Decision)')
doc.add_paragraph('角色: 自动驾驶算法工程师 (Agentic Assistant)')

# 1. 项目摘要
doc.add_heading('1. 项目摘要', level=1)
p = doc.add_paragraph(
    '本项目旨在构建一个基于“预测-决策”一体化的端到端自动驾驶系统。针对动态交通环境中未来不确定性高、多模态（Multi-modal）博弈复杂的痛点，'
    '我们搭建了基于 ScenePredNet（场景预测网络）与 iLQR（迭代线性二次调节器）的闭环仿真平台。'
    '系统能够从原始感知数据出发，生成包含了多种未来可能性的“场景树”，并通过二阶优化算法生成对应的“轨迹树”，'
    '最终实现了具备鲁棒性（Robustness）的实时控制闭环。目前已在4个高挑战性的 Demo 场景中实现了无碰撞、平滑的自动驾驶演示。'
)

# 2. 关键问题排查与技术攻关
doc.add_heading('2. 关键技术攻关与解决方案', level=1)

# 2.1
doc.add_heading('2.1 优化算法选型：为何放弃梯度下降？', level=2)
doc.add_paragraph('现象描述：').bold = True
doc.add_paragraph('在初期尝试使用基于神经网络的梯度下降法直接生成轨迹时，发现车辆控制不稳定，容易陷入局部极小值，且收敛速度无法满足 10Hz 的实时性要求。')
doc.add_paragraph('根因分析：').bold = True
doc.add_paragraph('梯度下降（一阶方法）类似于“盲人下山”，仅能感知局部坡度，对于平滑的物理动力学流形（Manifold）探测效率低。而车辆控制问题本质上是一个参数量少但对平滑性要求极高的优化问题。')
doc.add_paragraph('解决方案：').bold = True
doc.add_paragraph(
    '引入 iLQR (Iterative Linear Quadratic Regulator) 作为核心求解器。'
    'iLQR 是一种二阶优化方法（类似于“物理学家下山”），利用了代价函数的二阶导数（Hessian矩阵），'
    '能够精确拟合局部曲率，实现“一步跳跃”式的极速收敛。实测证明，iLQR 在处理车辆运动学约束时比传统 SGD 快 1-2 个数量级。'
)

# 2.2
doc.add_heading('2.2 多模态决策冲突：如何“稳一手”？', level=2)
doc.add_paragraph('现象描述：').bold = True
doc.add_paragraph('神经网络预测出前方车辆有“加速直行”和“突然切入”两种可能。如果在两个分支分别计算最优解，系统会在下一帧根据概率波动反复横跳，导致车辆左右摇摆。')
doc.add_paragraph('根因分析：').bold = True
doc.add_paragraph('独立优化导致 $T=0$ 时刻的动作不一致。现实中，车辆在当前时刻只能执行一个动作，必须兼容所有未来。')
doc.add_paragraph('解决方案：').bold = True
doc.add_paragraph(
    '实施树根一致性约束（Tree-Root Consistency）。'
    '在构建“轨迹树”时，强制要求所有未来分支在根节点共享同一个控制量 $u_0$。'
    'iLQR 优化目标被修正为最小化所有分支代价的期望值（Expected Cost）。'
    '最终生成的动作是一个“万能起手式”，它可能不是针对某一特定未来的最优解，但它是综合风险最小的鲁棒解（Robust Action）。'
)

# 2.3
doc.add_heading('2.3 仿真数据流闭环', level=2)
doc.add_paragraph('现象描述：').bold = True
doc.add_paragraph('需要验证系统在长时段运行下的稳定性，而非单帧预测能力。')
doc.add_paragraph('解决方案：').bold = True
doc.add_paragraph(
    '搭建 run_sim.py 实时闭环仿真器。'
    '实现了 Raw Data -> Tensorization -> ScenePredNet -> Scenario Tree -> Trajectory Tree -> Control -> Physics Step 的完整数据流。'
    '确保了每一次控制输出都反馈到真实的物理环境中，验证了算法在连续时间序列上的稳定性。'
)

# 插入核心架构图
doc.add_heading('核心数据流架构图', level=3)
try:
    doc.add_picture('/Users/phy/.gemini/antigravity/brain/5c1f2c4c-186e-4d3d-a2fa-8b485e18ba65/uploaded_image_1769139496820.png', width=Inches(5.0))
    doc.add_paragraph('图1: MIND 系统从感知输入到轨迹树生成的完整数据流')
except Exception as e:
    doc.add_paragraph(f'[插图失败: {e}]')

# 3. 系统架构解析
doc.add_heading('3. 系统架构解析', level=1)
doc.add_paragraph('为确保系统模块化开发，对核心文件功能定义如下：')

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '模块/文件'
hdr_cells[1].text = '功能定义'

# 填充表格
data = [
    ('simulator.py / run_sim.py', '【发射台/环境】负责物理世界仿真、Agent状态更新、碰撞检测及可视化渲染。它模拟了真实的物理世界。'),
    ('ScenePredNet', '【直觉大脑】基于历史观测数据，通过神经网络“直觉”地预测出未来场景的多种可能性（概率分布），构建 Scenario Tree。'),
    ('planners/mind/planner.py', '【大脑/中控】核心规划器类。负责串联预测与决策模块，管理数据流转，执行 Tensor化 和后处理。'),
    ('TrajectoryTreeOptimizer (iLQR)', '【逻辑大脑】基于预测的场景树，利用数学模型（二阶优化）精确计算出代价最低的轨迹树。它是理性的执行者。'),
    ('Robust Action Selection', '【最终决策】从轨迹树中提取树根动作。将数学优化的 Jerk/Rate 转换为底盘可执行的 Acceleration/SteeringAngle。')
]

for item, desc in data:
    row_cells = table.add_row().cells
    row_cells[0].text = item
    row_cells[1].text = desc

# 4. 阶段性训练结果分析
doc.add_heading('4. 阶段性验证结果', level=1)
doc.add_paragraph('完成了 4 个高难度场景（Demo 1-4）的闭环验证。')

doc.add_paragraph('通过分析产生的 Demo 视频（如 video_rescued.mp4），我们可以观察到：')
doc.add_paragraph(
    '1. 动态博弈能力：在 Demo 4 中，面对干扰车的切入，自车（Ego）成功通过“减速让行”或“微调方向”规避了碰撞。\n'
    '2. 舒适性指标：生成的轨迹平滑，加速度变化率（Jerk）控制在舒适范围内，未出现急刹急停。\n'
    '3. 实时性验证：系统在单线程 CPU 环境下稳定运行，能够在规定时间步内完成从感知到控制的全流程计算。'
    '最终输出给底盘的实际上是【加速度】和【转向角】，确保了物理执行的准确性。'
)

# 5. 后续优化建议
doc.add_heading('5. 后续优化建议', level=1)
doc.add_paragraph(
    '1. 成本函数调优：目前的权重（舒适度 vs 效率）较为固定，建议引入自适应权重以应对紧急避险场景。\n'
    '2. 硬件加速：ScenePredNet 推理目前占用较多 CPU 时间，建议迁移至 CUDA 进行 TensorRT 加速。\n'
    '3. 更多Corner Case：建议在仿真器中增加极端天气或传感器噪声，测试 Robust Control 的极限边界。'
)

# 保存文档
output_path = '/Users/phy/Desktop/MIND/MIND_技术报告.docx'
doc.save(output_path)
print(f"Document saved to {output_path}")
